from slurm.types import Category, Subcategory, Tossup

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from html.parser import HTMLParser
from typing import List, Optional, Self


class FileWriter:
    """ Interface for writing tossups to a file. """

    def write_tossups(
        self: Self,
        packet_name: str,
        tossups: List[Tossup],
        basename: Optional[str]=None
    ) -> None:
        """ Writes the tossups to a document with the given basename and assigns the given name to the packet. """
        pass


class MicrosoftWordWriter(FileWriter):
    """ File writer which writes tossups to a Microsoft Word document. """

    def write_tossups(
        self: Self,
        packet_name: str,
        tossups: List[Tossup],
        basename: Optional[str]=None
    ) -> None:
        document = Document()

        self.set_styles(document)

        self.write_heading(document, packet_name)

        for tossup in tossups:
            self.write_question(document, tossup.question)
            self.write_answerline(document, tossup.formatted_answer)
            self.write_category_and_subcategory(document, tossup.category, tossup.subcategory)
            self.write_empty_line(document)

        document.save(f"{basename if basename else packet_name}.docx")

    def set_styles(self: Self, document: Document) -> None:
        document.styles["Normal"].font.name = "Arial"

        style = document.styles.add_style('Packet Name', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style.font.size = Pt(18)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_after = Pt(16)

    def write_heading(self: Self, document: Document, heading: str) -> None:
        document.add_paragraph(heading, style="Packet Name")

    def write_question(self: Self, document: Document, question: str) -> None:
        document.add_paragraph(question, style="List Number")

    def write_answerline(self: Self, document: Document, answerline: str) -> None:
        paragraph = document.add_paragraph("ANSWER: ", style="List Continue")
        parser = self.MicrosoftWordAnswerLineParser(paragraph)
        parser.feed(answerline)

    def write_category_and_subcategory(self: Self, document: Document, category: Category, subcategory: Subcategory):
        document.add_paragraph(f"<{category} / {subcategory}>", style="List Continue")

    def write_empty_line(self: Self, document: Document):
        document.add_paragraph("", style="List Continue")

    class MicrosoftWordAnswerLineParser(HTMLParser):
        def __init__(self: Self, paragraph: Paragraph):
            self.paragraph = paragraph
            self.bold = self.underline = False
            super().__init__()

        def handle_starttag(self: Self, tag: str, attrs):
            if tag == "u":
                self.underline = True
            if tag == "b":
                self.bold = True

        def handle_endtag(self: Self, tag: str):
            if tag == "u":
                self.underline = False
            if tag == "b":
                self.bold = False

        def handle_data(self: Self, data: str):
            run = self.paragraph.add_run(data)
            run.underline = self.underline
            run.bold = self.bold